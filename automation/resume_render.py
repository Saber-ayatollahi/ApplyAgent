#!/usr/bin/env python3
"""
resume_render.py — THE canonical resume renderer for Saber Ayatollahi's job search.

ONE template. ONE script. The branded layout/style lives entirely in this file
(no .docx template on disk to get locked, edited, or drift out of sync). It
consumes a structured *resume content* JSON — the per-job tailored content an AI
agent produces from the Master Repo + a job description — and emits a polished,
ATS-aware .docx (and .pdf) that matches the approved house style.

    Pipeline:  Master Repo (+ JD)  --[AI tailoring, see docs/resume_agent_instructions.md]-->
               resume_content.json  --[this script]-->  applications/<date>_<company>_<role>/  (.docx + .pdf)

DESIGN PRINCIPLES (from the 2026 resume research — see docs/resume_agent_instructions.md):
  * SINGLE COLUMN, NO TABLES. Tables / multi-column grids are the #1 cause of ATS
    parse failures (scanners read across columns and drop/scramble content). Every
    section here is linear paragraphs so Workday/Greenhouse/Taleo/iCIMS parse it
    top-to-bottom, left-to-right, exactly as written.
  * Standard, parseable section headings (SUMMARY / CORE SKILLS / EXPERIENCE /
    EDUCATION). A thin rule under each header is a paragraph border (ATS-safe), not
    a table.
  * Real font name preserved (Avenir Next LT Pro) so it matches the candidate's
    installed brand in Word; the PDF renderer substitutes a close free geometric
    sans (Montserrat) when Avenir is not installed.

Usage:
    python resume_render.py --content automation/resume_data/<job>.json --check-pages
    python resume_render.py --schema      # print the JSON content schema
    python resume_render.py --example     # print an example content JSON

Dependencies: python-docx. PDF + page-check use libreoffice/soffice + pypdf if present.
"""
from __future__ import annotations
import argparse
import json
import math
import re
import shutil
import subprocess
import sys
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.stderr.write("python-docx is required:  pip install python-docx\n")
    raise

# ---------------------------------------------------------------------------
# HOUSE STYLE — single source of truth for the template's look.
# ---------------------------------------------------------------------------
FONT        = "Avenir Next LT Pro"        # headings, body, bullets
FONT_LIGHT  = "Avenir Next LT Pro Light"  # summary paragraph
INK         = "1A1A1A"                     # near-black (softer than pure #000)
RULE        = "BFBFBF"                     # section-header underline grey

MARGIN_LR   = 0.6
MARGIN_TOP  = 0.5
MARGIN_BOT  = 0.5

SZ_NAME     = 19
SZ_CRED     = 11
SZ_CONTACT  = 9.5
SZ_SECTION  = 11.5
SZ_SUMMARY  = 10
SZ_SKILL    = 10
SZ_EMPLOYER = 11.5
SZ_ROLE     = 10.5
SZ_SUBHEAD  = 10.5
SZ_BULLET   = 10

SECTION_SUMMARY    = "SUMMARY"
SECTION_SKILLS     = "CORE SKILLS"
SECTION_EXPERIENCE = "EXPERIENCE"
SECTION_EDUCATION  = "EDUCATION & CREDENTIALS"

DOT = "  •  "   # mid-dot separator for the skills line / contact line

# Two-page length budget (rough chars/line over the usable width).
CHARS_PER_LINE = 105
TWO_PAGE_LINE_BUDGET = 92

# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------
def _run(p, text, *, size=None, bold=False, italic=False, font=FONT, color=INK, caps=False):
    r = p.add_run(text)
    r.font.name = font
    rpr = r._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(a), font)
    if size is not None:
        r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = RGBColor.from_string(color)
    if caps:
        cs = OxmlElement('w:caps'); cs.set(qn('w:val'), 'true'); rpr.append(cs)
        sp = OxmlElement('w:spacing'); sp.set(qn('w:val'), '20'); rpr.append(sp)  # subtle tracking
    return r

def _para(doc, *, align=None, before=0, after=0, line=None,
          left=0.0, hanging=0.0, keep_next=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line
    if left:
        pf.left_indent = Inches(left)
    if hanging:
        pf.first_line_indent = Inches(-hanging)
        if left:
            pf.tab_stops.add_tab_stop(Inches(left))  # align bullet text + wraps
    if keep_next:
        pf.keep_with_next = True
    return p

def _bottom_rule(p, color=RULE, sz=6, space=3):
    pPr = p._p.get_or_add_pPr()
    pbdr = pPr.find(qn('w:pBdr'))
    if pbdr is None:
        pbdr = OxmlElement('w:pBdr'); pPr.append(pbdr)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), str(space)); bottom.set(qn('w:color'), color)
    pbdr.append(bottom)

# ---------------------------------------------------------------------------
# Section builders (single column, no tables)
# ---------------------------------------------------------------------------
def _build_header(doc, c):
    p = _para(doc, after=1, line=1.0)
    _run(p, c["name"], size=SZ_NAME, bold=True)
    if c.get("credentials"):
        _run(p, ", " + c["credentials"], size=SZ_CRED, bold=False)
    # one linear contact line — fully ATS-parseable
    bits = [c.get("location"), c.get("phone"), c.get("email")]
    if c.get("linkedin"):
        bits.append(c["linkedin"].split("://", 1)[-1].rstrip("/"))
    contact = DOT.join(b for b in bits if b)
    pc = _para(doc, after=4, line=1.0)
    _run(pc, contact, size=SZ_CONTACT)

def _section(doc, title):
    p = _para(doc, before=7, after=3, line=1.0, keep_next=True)
    _run(p, title, size=SZ_SECTION, bold=True, caps=True)
    _bottom_rule(p)

def _build_summary(doc, summary):
    p = _para(doc, after=2, line=1.08, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    _run(p, summary, size=SZ_SUMMARY, font=FONT_LIGHT)

def _build_skills(doc, skills):
    # one linear paragraph of dot-separated terms (ATS reads it straight through)
    p = _para(doc, after=2, line=1.12)
    last = len(skills) - 1
    for i, s in enumerate(skills):
        _run(p, s, size=SZ_SKILL)
        if i != last:
            _run(p, DOT, size=SZ_SKILL, color=RULE)

def _build_experience(doc, experience):
    for emp in experience:
        pe = _para(doc, before=5, after=0, line=1.0, keep_next=True)
        _run(pe, emp["employer"], size=SZ_EMPLOYER, bold=True, italic=True)
        for role in emp["roles"]:
            pr = _para(doc, after=2, line=1.0, keep_next=True)
            _run(pr, role["title"], size=SZ_ROLE, bold=True, italic=True)
            if role.get("location_date"):
                _run(pr, "  |  " + role["location_date"], size=SZ_ROLE, italic=True)
            for sec in role.get("sections", []):
                if sec.get("heading"):
                    ph = _para(doc, before=3, after=1, line=1.0, keep_next=True)
                    _run(ph, sec["heading"], size=SZ_SUBHEAD, bold=True)
                for b in sec.get("bullets", []):
                    pb = _para(doc, after=2, line=1.08, left=0.18, hanging=0.18,
                               align=WD_ALIGN_PARAGRAPH.JUSTIFY)
                    _run(pb, "•\t" + b, size=SZ_BULLET)

def _build_education(doc, education):
    for e in education:
        pb = _para(doc, after=1, line=1.08, left=0.18, hanging=0.18)
        _run(pb, "•\t" + e, size=SZ_BULLET)

# ---------------------------------------------------------------------------
# Validation / length / keyword self-check
# ---------------------------------------------------------------------------
def validate(content):
    """Validate that `content` has every field the render path dereferences.

    Contract: if validate() passes, render()/estimate_lines() will not raise a
    bare KeyError — every []-indexed field below is required here. (Previously
    only the top-level keys + employer/roles were checked, so malformed-but-
    passing content crashed mid-render with a confusing traceback.)"""
    errs = []
    for key in ("contact", "summary", "core_skills", "experience", "education"):
        if key not in content:
            errs.append(f"missing required key: {key}")
    if "name" not in content.get("contact", {}):
        errs.append("contact missing required 'name'")
    for emp in content.get("experience", []):
        if "employer" not in emp or "roles" not in emp:
            errs.append(f"experience entry missing employer/roles: {emp.get('employer','?')}")
            continue
        for role in emp["roles"]:
            if "title" not in role:
                errs.append(f"role under '{emp['employer']}' missing required 'title'")
    if errs:
        raise ValueError("Content validation failed:\n  - " + "\n  - ".join(errs))

def _all_text(content):
    parts = [content.get("summary", "")]
    parts += content.get("core_skills", [])
    for emp in content.get("experience", []):
        parts.append(emp["employer"])
        for role in emp["roles"]:
            parts.append(role.get("title", "") + " " + role.get("location_date", ""))
            for sec in role.get("sections", []):
                if sec.get("heading"):
                    parts.append(sec["heading"])
                parts += sec.get("bullets", [])
    parts += content.get("education", [])
    return "\n".join(parts)

def estimate_lines(content):
    total = 5
    total += math.ceil(len(content.get("summary", "")) / CHARS_PER_LINE) + 1
    skills_len = sum(len(s) + 3 for s in content.get("core_skills", []))
    total += math.ceil(skills_len / CHARS_PER_LINE) + 1
    for emp in content.get("experience", []):
        total += 1
        for role in emp["roles"]:
            total += 1
            for sec in role.get("sections", []):
                if sec.get("heading"):
                    total += 1
                for b in sec.get("bullets", []):
                    total += math.ceil((len(b) + 3) / CHARS_PER_LINE)
    total += len(content.get("education", [])) + 1
    return total

def _kw_present(keyword, text):
    """True if `keyword` appears in `text` (both lowercased) as a whole token,
    not embedded in a larger word — so 'SQL' does NOT match inside 'PostgreSQL'
    and 'Excel' does NOT match inside 'excellent'. Boundaries are non-word
    characters; this keeps multi-word and punctuated keywords ('C++', 'value at
    risk') working."""
    k = keyword.lower().strip()
    if not k:
        return False
    # Lowercase `text` here too so the whole-token match is genuinely
    # case-insensitive regardless of caller (keyword_report already passes
    # lowercased text; .lower() is idempotent so this is free for it).
    return re.search(r"(?<!\w)" + re.escape(k) + r"(?!\w)",
                     text.lower()) is not None

def keyword_report(content):
    text = _all_text(content).lower()
    kws = content.get("target", {}).get("jd_keywords", [])
    missing = [k for k in kws if not _kw_present(k, text)]
    return kws, missing

# ---------------------------------------------------------------------------
# Render
# ---------------------------------------------------------------------------
def render(content, out_path):
    validate(content)
    doc = Document()
    for sec in doc.sections:
        sec.left_margin = Inches(MARGIN_LR); sec.right_margin = Inches(MARGIN_LR)
        sec.top_margin = Inches(MARGIN_TOP); sec.bottom_margin = Inches(MARGIN_BOT)
    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal.font.size = Pt(SZ_BULLET)

    _build_header(doc, content["contact"])
    _section(doc, SECTION_SUMMARY);    _build_summary(doc, content["summary"])
    _section(doc, SECTION_SKILLS);     _build_skills(doc, content["core_skills"])
    label = content.get("target", {}).get("experience_label", SECTION_EXPERIENCE)
    _section(doc, label);              _build_experience(doc, content["experience"])
    _section(doc, SECTION_EDUCATION);  _build_education(doc, content["education"])

    doc.save(out_path)
    return out_path


def _clean_cover_body(text):
    """Strip markdown noise from the LLM cover text into clean letter
    paragraphs: drop headings, horizontal rules, bold markers, leading
    bullets, and any contact-header or honest-gaps block (the .docx supplies
    the branded header; gaps live in notes, not the letter)."""
    lines = []
    for ln in (text or "").splitlines():
        st = ln.strip()
        if st.startswith("#") or st in ("---", "***", "___"):
            continue
        if (re.match(r"^\*{0,2}Saber Ayatollahi", st)
                or st.lower().startswith("subject:")
                or "saber.ayatollahi@gmail.com" in st
                or "linkedin.com/in/sayatollahi" in st):
            continue
        if st.startswith("*") and st.endswith("*") and "gap" in st.lower():
            continue
        s = re.sub(r"\*\*(.+?)\*\*", r"\1", ln.rstrip())  # bold
        s = re.sub(r"^\s*[-*]\s+", "", s)                  # leading bullets
        lines.append(s)
    return "\n".join(lines).strip()


def render_cover(contact, target, body, out_path):
    """Render a branded cover-letter .docx — same header + typography as the
    resume, then a dated letter body. `body` is the letter text (salutation →
    paragraphs → signature); markdown is stripped to clean paragraphs. Pairs
    with to_pdf() for a matching PDF."""
    doc = Document()
    for sec in doc.sections:
        sec.left_margin = Inches(MARGIN_LR); sec.right_margin = Inches(MARGIN_LR)
        sec.top_margin = Inches(MARGIN_TOP); sec.bottom_margin = Inches(MARGIN_BOT)
    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal.font.size = Pt(SZ_BULLET)

    _build_header(doc, contact or {})
    doc.add_paragraph(date.today().strftime("%B %d, %Y"))
    company = (target or {}).get("company")
    role = (target or {}).get("role")
    if company:
        doc.add_paragraph(f"Re: {role} — {company}" if role else f"Re: {company}")
    for para in _clean_cover_body(body).split("\n\n"):
        para = para.strip()
        if para:
            doc.add_paragraph(para)
    doc.save(out_path)
    return out_path

# ---------------------------------------------------------------------------
# Output bundling — folder per job, both .docx and .pdf
# ---------------------------------------------------------------------------
def _slug(text):
    """Filesystem-safe slug for folder/file names. Drops apostrophes (so
    "Moody's" -> "Moodys", not "Moody-s") and transliterates accented letters
    to ASCII (NFKD fold, so "Societe Generale" with accents degrades to
    "Societe-Generale") before collapsing any remaining non-alphanumerics."""
    import unicodedata
    s = unicodedata.normalize("NFKD", str(text))
    s = s.encode("ascii", "ignore").decode("ascii")   # drop accents
    s = s.replace("'", "").replace("’", "")        # drop straight + curly apostrophes
    out = re.sub(r"[^A-Za-z0-9]+", "-", s).strip("-")
    # An all-non-ASCII name (e.g. CJK) folds to ''; never return empty, or the
    # folder/file name would collapse to "<date>__<role>".
    return out or "x"

# Max seconds to wait for a single LibreOffice headless conversion before giving
# up. A hung/locked soffice (shared-profile lock, recovery dialog) would
# otherwise block the whole render forever; we'd rather skip the PDF.
PDF_CONVERT_TIMEOUT = 120

def _to_pdf_via_word(docx_path, out_dir):
    """Fallback PDF path for machines with MS Word but no LibreOffice
    (e.g. Windows). Uses docx2pdf, which drives Word via COM. Slower
    (~10s, opens Word) and Windows/Word-only; returns Path or None."""
    try:
        from docx2pdf import convert  # type: ignore
    except Exception:
        return None
    dest = Path(out_dir) / (Path(docx_path).stem + ".pdf")
    try:
        convert(str(docx_path), str(dest))
        return dest if dest.exists() else None
    except Exception as e:  # noqa: BLE001
        print(f"  WARNING: Word/docx2pdf PDF conversion failed: {e}")
        return None


def to_pdf(docx_path, out_dir):
    """Convert .docx -> .pdf via libreoffice/soffice in a private temp dir so
    LibreOffice lock/temp droppings never land in the deliverable folder; only
    the finished PDF is moved into out_dir. Returns Path or None.

    The conversion is bounded by PDF_CONVERT_TIMEOUT so a hung soffice can't
    wedge the pipeline, and the finished PDF is placed via an atomic os.replace
    so an interrupted run never leaves a half-written or silently-stale
    canonical PDF. If the canonical path is genuinely locked (open in a viewer
    on the host) we fall back to a '-new.pdf' copy and warn LOUDLY that the
    canonical file is now stale.

    When LibreOffice isn't on PATH, falls back to MS Word (docx2pdf) before
    giving up — so a Word-equipped Windows box still gets PDFs."""
    soffice = shutil.which("libreoffice") or shutil.which("soffice")
    if not soffice:
        pdf = _to_pdf_via_word(docx_path, out_dir)
        if pdf:
            return pdf
        print("  WARNING: PDF skipped (no libreoffice/soffice on PATH, and "
              "Word/docx2pdf unavailable).")
        return None
    import os
    import tempfile
    stem = Path(docx_path).stem
    # Isolate LibreOffice's per-user profile to a throwaway dir: the shared
    # profile lock is the most common cause of a headless-soffice hang.
    with tempfile.TemporaryDirectory() as td:
        profile = Path(td) / "lo_profile"
        # as_uri() emits the schema-correct file:///C:/... on Windows and
        # file:///tmp/... on POSIX. A hand-built "file://" + path would put the
        # drive letter in the URI authority on Windows (file://C:/…), which LO
        # ignores — falling back to the shared profile we're trying to avoid.
        try:
            proc = subprocess.run(
                [soffice, f"-env:UserInstallation={profile.as_uri()}",
                 "--headless", "--convert-to", "pdf", "--outdir", td, str(docx_path)],
                check=False, capture_output=True, timeout=PDF_CONVERT_TIMEOUT)
        except subprocess.TimeoutExpired:
            print(f"  WARNING: PDF conversion timed out after {PDF_CONVERT_TIMEOUT}s "
                  f"(LibreOffice hung/locked); PDF skipped.")
            return None
        tmp_pdf = Path(td) / (stem + ".pdf")
        if not tmp_pdf.exists():
            # soffice WAS found and ran, but produced no PDF (corrupt docx,
            # soffice crash, unsupported feature). Surface the real cause —
            # never report it as "not found on PATH". stderr/stdout are
            # captured above; decode best-effort and trim to a readable tail.
            def _tail(stream: bytes) -> str:
                return (stream or b"").decode("utf-8", "replace").strip()[-300:]
            detail = _tail(proc.stderr) or _tail(proc.stdout) or "no diagnostic output"
            print(f"  WARNING: PDF conversion failed (soffice exit {proc.returncode}, "
                  f"no PDF produced); PDF skipped. soffice said: {detail}")
            return None
        dest = Path(out_dir) / (stem + ".pdf")
        # Stage the result in the destination dir so os.replace is atomic
        # (same filesystem) and never leaves a partial canonical file.
        staged = Path(out_dir) / (stem + ".pdf.tmp")
        shutil.copyfile(tmp_pdf, staged)
        try:
            os.replace(staged, dest)   # atomic overwrite of the canonical name
            return dest
        except PermissionError:
            # canonical path is locked (e.g. open in a PDF viewer) — we cannot
            # overwrite or delete an open file on Windows; keep the fresh render
            # under -new.pdf and make it unmistakable that dest is now STALE.
            alt = Path(out_dir) / (stem + "-new.pdf")
            try:
                os.replace(staged, alt)
            except OSError:
                # even -new.pdf is locked — never leave the staging temp behind
                # or crash the pipeline; drop it and report.
                try:
                    staged.unlink()
                except OSError:
                    pass
                print(f"  WARNING: both {dest.name} and {alt.name} are open/locked; "
                      f"PDF not written. Close the open PDF(s) and re-run.")
                return None
            print(f"  WARNING: {dest.name} is open/locked and is now STALE; wrote "
                  f"{alt.name} with the current render. Close the open PDF, delete "
                  f"the stale {dest.name}, and rename {alt.name} (or re-run).")
            return alt

def bundle(content, base_dir, make_pdf=True, on_date=None):
    tgt = content.get("target", {})
    company = tgt.get("company", "Company")
    role = tgt.get("role", "Role")
    d = on_date or date.today().isoformat()
    folder = Path(base_dir) / f"{d}_{_slug(company)}_{_slug(role)}"
    folder.mkdir(parents=True, exist_ok=True)
    stem = f"Saber_Ayatollahi_{_slug(company)}_{_slug(role)}"
    docx_path = folder / (stem + ".docx")
    render(content, str(docx_path))
    pdf_path = to_pdf(docx_path, folder) if make_pdf else None
    return folder, docx_path, pdf_path

# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
# The actual content-JSON schema (what --schema prints). This is the key shape
# an AI agent must produce — NOT the module docstring (which is design prose).
SCHEMA = """resume_content.json — required key shape:

{
  "contact": {                 # REQUIRED. 'name' is required; rest optional.
    "name":        str,        #   REQUIRED
    "credentials": str,        #   optional, e.g. "CFA, MSc"
    "location":    str,        #   optional
    "phone":       str,        #   optional
    "email":       str,        #   optional
    "linkedin":    str         #   optional, full URL
  },
  "target": {                  # optional, but drives folder name + ATS check
    "company":          str,   #   used in folder/file name
    "role":             str,   #   used in folder/file name
    "experience_label": str,   #   optional, overrides the EXPERIENCE heading
    "jd_keywords":      [str]  #   ATS self-check (whole-token match)
  },
  "summary":     str,          # REQUIRED. One paragraph, 60-85 words.
  "core_skills": [str],        # REQUIRED. Dot-separated on one ATS-safe line.
  "experience": [              # REQUIRED. List of employers.
    {
      "employer": str,         #   REQUIRED
      "roles": [               #   REQUIRED
        {
          "title":         str,   # REQUIRED
          "location_date": str,   # optional, e.g. "Toronto, May 2023 - Present"
          "sections": [           # optional
            {"heading": str|null, # optional sub-heading
             "bullets": [str]}    # bullet lines
          ]
        }
      ]
    }
  ],
  "education": [str]           # REQUIRED. One string per credential line.
}

Run --example for a concrete, renderable sample of this shape.
"""

EXAMPLE = {
    "contact": {"name": "Saber Ayatollahi", "credentials": "CFA, MSc",
                "location": "Toronto, Canada", "phone": "+1 (416) 856-1276",
                "email": "saber.ayatollahi@gmail.com",
                "linkedin": "https://www.linkedin.com/in/sayatollahi/"},
    "target": {"company": "Example Bank", "role": "Director, Risk",
               "jd_keywords": ["market risk", "Python", "SQL"]},
    "summary": "One tight paragraph, 60-85 words, evidence-backed, opens with target title + years.",
    "core_skills": ["Skill A", "Skill B", "Skill C", "Skill D", "Skill E", "Skill F"],
    "experience": [
        {"employer": "Employer One",
         "roles": [
            {"title": "Senior Title", "location_date": "Toronto, May 2023 - Present",
             "sections": [
                {"heading": "Theme Heading", "bullets": ["Action verb + quantified result bullet."]},
                {"heading": None, "bullets": ["A plain bullet with no sub-heading."]}
             ]}
         ]}
    ],
    "education": ["Credential - Institution, Year"]
}

def main():
    ap = argparse.ArgumentParser(description="Render a tailored resume into a dated, job-titled folder as .docx + .pdf.")
    ap.add_argument("--content", help="path to resume content JSON")
    ap.add_argument("--bundle-base", default="applications", help="base dir for per-job folders (default: applications)")
    ap.add_argument("--out", help="LEGACY: render a single .docx to this exact path instead of a bundle")
    ap.add_argument("--no-pdf", action="store_true", help="skip PDF generation")
    ap.add_argument("--date", help="override folder date (YYYY-MM-DD); default today")
    ap.add_argument("--schema", action="store_true", help="print the content schema and exit")
    ap.add_argument("--example", action="store_true", help="print an example content JSON and exit")
    ap.add_argument("--check-pages", action="store_true", help="report PDF page count")
    args = ap.parse_args()

    if args.schema:
        print(SCHEMA); return 0
    if args.example:
        print(json.dumps(EXAMPLE, indent=2, ensure_ascii=False)); return 0
    if not args.content:
        ap.error("--content is required (or use --schema/--example)")

    content = json.loads(Path(args.content).read_text(encoding="utf-8"))

    pdf_path = None
    if args.out:
        out = render(content, args.out)
        print(f"rendered {out}")
        if not args.no_pdf:
            pdf_path = to_pdf(out, Path(out).resolve().parent)
            if pdf_path:
                print(f"pdf      {pdf_path}")
            # else: to_pdf() already printed the specific cause (not on PATH /
            # timeout / conversion failed / file locked) — don't override it.
    else:
        folder, docx_path, pdf_path = bundle(content, args.bundle_base,
                                             make_pdf=not args.no_pdf, on_date=args.date)
        print(f"folder   {folder}")
        print(f"docx     {docx_path.name}")
        if pdf_path:
            print(f"pdf      {pdf_path.name}")
        # elif not args.no_pdf: to_pdf() (called inside bundle()) already
        # printed the specific cause — don't override it with a guess.

    lines = estimate_lines(content)
    flag = "OK" if lines <= TWO_PAGE_LINE_BUDGET else "OVER 2-PAGE BUDGET"
    print(f"  length estimate: ~{lines} lines (budget {TWO_PAGE_LINE_BUDGET}) [{flag}]")

    kws, missing = keyword_report(content)
    if kws:
        if missing:
            print(f"  ATS keywords MISSING ({len(missing)}/{len(kws)}): {', '.join(missing)}")
        else:
            print(f"  all {len(kws)} target ATS keywords present")

    if args.check_pages:
        try:
            from pypdf import PdfReader
            if pdf_path and Path(pdf_path).exists():
                n = len(PdfReader(str(pdf_path)).pages)
                print(f"  pages: {n} [{'OK' if n <= 2 else 'MORE THAN 2 PAGES'}]")
            else:
                print("  (page check needs a generated PDF)")
        except Exception as e:
            print(f"  (page check unavailable: {e})")
    return 0

if __name__ == "__main__":
    raise SystemExit(main())
