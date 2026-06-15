"""Tests for automation/resume_render.py — the canonical resume renderer whose
.docx/.pdf output is sent to real employers. The module shipped with zero
coverage despite a ~30-file test culture; this exercises the pure, deterministic
surfaces (validate / render / estimate_lines / keyword_report / _slug / bundle /
to_pdf graceful-degrade) so a silent regression can't quietly degrade a
document an employer reads.

Pure + network-free: render() needs python-docx (a declared dep); to_pdf() is
tested only on the no-soffice path (monkeypatched) so it never shells out."""
from __future__ import annotations
import json
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

import resume_render as rr  # type: ignore
from docx import Document   # type: ignore

RBC_JSON = ROOT / "resume_data" / "rbc_global_risk_analytics.json"


def _minimal():
    return {
        "contact": {"name": "Test Person"},
        "summary": "A short summary.",
        "core_skills": ["Skill A", "Skill B"],
        "experience": [
            {"employer": "Employer One",
             "roles": [{"title": "Senior Title",
                        "sections": [{"heading": None, "bullets": ["A bullet."]}]}]}
        ],
        "education": ["Credential - Institution, Year"],
    }


# --- validate() -------------------------------------------------------------
def test_validate_accepts_minimal():
    rr.validate(_minimal())  # must not raise


def test_validate_accepts_example_and_rbc():
    rr.validate(rr.EXAMPLE)
    rr.validate(json.loads(RBC_JSON.read_text(encoding="utf-8")))


def test_validate_rejects_missing_toplevel_key():
    bad = _minimal(); del bad["summary"]
    _assert_raises(lambda: rr.validate(bad), "missing summary")


def test_validate_rejects_missing_contact_name():
    bad = _minimal(); bad["contact"] = {"email": "a@b.c"}
    _assert_raises(lambda: rr.validate(bad), "contact without name")


def test_validate_rejects_role_without_title():
    bad = _minimal(); bad["experience"][0]["roles"][0] = {"sections": []}
    _assert_raises(lambda: rr.validate(bad), "role without title")


def test_validate_rejects_experience_without_roles():
    bad = _minimal(); bad["experience"] = [{"employer": "E"}]
    _assert_raises(lambda: rr.validate(bad), "experience without roles")


def test_validate_contract_render_does_not_keyerror():
    """If validate() passes, render() must not raise a bare KeyError."""
    import tempfile, os
    c = _minimal()
    rr.validate(c)
    with tempfile.TemporaryDirectory() as td:
        rr.render(c, os.path.join(td, "x.docx"))  # must not KeyError


# --- render() round-trip ----------------------------------------------------
def test_render_roundtrip_no_tables_and_content_present():
    import tempfile, os
    c = json.loads(RBC_JSON.read_text(encoding="utf-8"))
    with tempfile.TemporaryDirectory() as td:
        path = os.path.join(td, "rbc.docx")
        rr.render(c, path)
        doc = Document(path)
        # The #1 stated ATS rule: single column, NO TABLES.
        assert len(doc.tables) == 0, "renderer emitted a table (breaks ATS parsing)"
        text = "\n".join(p.text for p in doc.paragraphs)
        assert c["contact"]["name"] in text
        for skill in c["core_skills"]:
            assert skill in text, f"missing skill: {skill}"
        # every bullet present
        for emp in c["experience"]:
            for role in emp["roles"]:
                for sec in role.get("sections", []):
                    for b in sec.get("bullets", []):
                        assert b in text, f"missing bullet: {b[:40]}"


# --- estimate_lines() -------------------------------------------------------
def test_estimate_lines_monotonic():
    c = _minimal()
    base = rr.estimate_lines(c)
    c["experience"][0]["roles"][0]["sections"][0]["bullets"].append("Another bullet line.")
    assert rr.estimate_lines(c) > base


def test_estimate_lines_rbc_under_budget():
    c = json.loads(RBC_JSON.read_text(encoding="utf-8"))
    assert rr.estimate_lines(c) <= rr.TWO_PAGE_LINE_BUDGET


# --- keyword_report() (word-boundary matching) ------------------------------
def test_keyword_report_word_boundary_no_false_positive():
    c = _minimal()
    c["summary"] = "I am excellent with PostgreSQL and noSQL stores."
    c["target"] = {"jd_keywords": ["SQL", "Excel"]}
    _, missing = rr.keyword_report(c)
    assert "SQL" in missing, "'SQL' must NOT match inside PostgreSQL/noSQL"
    assert "Excel" in missing, "'Excel' must NOT match inside 'excellent'"


def test_keyword_report_present_and_punctuated():
    c = _minimal()
    c["summary"] = "We compute value at risk in Python and C++ daily."
    c["target"] = {"jd_keywords": ["value at risk", "Python", "C++"]}
    _, missing = rr.keyword_report(c)
    assert missing == [], f"expected all present, missing={missing}"


def test_keyword_report_rbc_all_present():
    c = json.loads(RBC_JSON.read_text(encoding="utf-8"))
    kws, missing = rr.keyword_report(c)
    assert kws and missing == [], f"RBC keywords should all be present, missing={missing}"


def test_keyword_report_absent_target_ok():
    kws, missing = rr.keyword_report(_minimal())  # no 'target' key
    assert kws == [] and missing == []


# --- _slug() ----------------------------------------------------------------
def test_slug_apostrophe_and_accents():
    assert rr._slug("Moody's Analytics") == "Moodys-Analytics"
    assert rr._slug("Société Générale") == "Societe-Generale"
    assert rr._slug("AT&T / Bell") == "AT-T-Bell"
    assert rr._slug("--Edge--Case--") == "Edge-Case"


# --- bundle() folder/file naming --------------------------------------------
def test_bundle_naming_and_no_pdf():
    import tempfile
    c = _minimal()
    c["target"] = {"company": "Moody's Analytics", "role": "Director, Risk"}
    with tempfile.TemporaryDirectory() as td:
        folder, docx_path, pdf_path = rr.bundle(c, td, make_pdf=False, on_date="2026-06-01")
        assert folder.name == "2026-06-01_Moodys-Analytics_Director-Risk"
        assert docx_path.name == "Saber_Ayatollahi_Moodys-Analytics_Director-Risk.docx"
        assert docx_path.exists()
        assert pdf_path is None


def test_bundle_bounds_long_title_for_max_path():
    """A verbose role title must not blow past Windows MAX_PATH (260) — the PDF
    converter silently fails there. bundle caps company[:20] + role[:40] for
    both the folder AND the filename, so the two stay consistent and short."""
    import tempfile
    c = _minimal()
    c["target"] = {"company": "Deloitte",
                   "role": "Manager/Senior Manager, Quantitative Market Risk "
                           "Models - Financial Engineering and Modeling"}
    with tempfile.TemporaryDirectory() as td:
        folder, docx_path, _pdf = rr.bundle(c, td, make_pdf=False,
                                            on_date="2026-06-14")
        # role slug in the folder is capped (no trailing dash from the cut)
        role_part = folder.name.split("_", 2)[-1]
        assert len(role_part) <= 40 and not role_part.endswith("-")
        # folder name and the resume stem use the SAME capped slugs
        assert docx_path.name == f"Saber_Ayatollahi_{folder.name.split('_', 1)[1]}.docx"
        # the full resume path stays comfortably under MAX_PATH
        assert len(str(docx_path.resolve())) < 200


# --- to_pdf() graceful no-soffice -------------------------------------------
def test_to_pdf_returns_none_without_soffice(monkeypatch=None):
    # No pytest dependency on monkeypatch fixture — patch shutil.which directly.
    import shutil
    orig = shutil.which
    shutil.which = lambda name: None
    try:
        assert rr.to_pdf("whatever.docx", ".") is None
    finally:
        shutil.which = orig


# --- helpers / runner -------------------------------------------------------
def _assert_raises(fn, label):
    try:
        fn()
    except ValueError:
        return
    raise AssertionError(f"{label}: expected ValueError, none raised")


def main() -> int:
    fns = [v for k, v in sorted(globals().items())
           if k.startswith("test_") and callable(v)]
    failed = 0
    for fn in fns:
        try:
            fn()
            print(f"  OK   {fn.__name__}")
        except Exception as e:  # noqa: BLE001
            failed += 1
            print(f"  FAIL {fn.__name__}: {e}")
    print(f"\n{len(fns) - failed}/{len(fns)} passed")
    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
